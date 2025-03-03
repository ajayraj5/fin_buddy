import pandas as pd # type: ignore
import time
from datetime import datetime
from bs4 import BeautifulSoup # type: ignore
import os
import re

from selenium import webdriver # type: ignore
from selenium.webdriver.chrome.options import Options # type: ignore
from selenium.webdriver.support.ui import WebDriverWait # type: ignore
from selenium.webdriver.support import expected_conditions as EC # type: ignore
from selenium.webdriver.common.by import By # type: ignore
from selenium.webdriver.common.keys import Keys # type: ignore
from selenium.webdriver.common.action_chains import ActionChains # type: ignore


import frappe
import openai  # type: ignore
import PyPDF2  # type: ignore
from PyPDF2.errors import PdfReadError
import io
from io import BytesIO
from docx import Document
from pathlib import Path
import json
from typing import Dict, Any, Optional

# from fin_buddy.events.ContentMaking import ContentMasker
from fin_buddy.events.DataMakingUtil import DataMaskingUtil
import fin_buddy.utils.dates as MakeDateValid
import tempfile


def setup_chrome_options(user_download_dir):
    """Setup Chrome options with the specified download directory"""
    
    # Create a unique temporary directory for user data
    unique_user_data_dir = os.path.join(tempfile.gettempdir(), f"chrome_user_data_{os.getpid()}")
    
    options = Options()
    options.add_argument("user-agent=Mozilla/5.0 (Windows NT 20.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36")
    options.add_argument("--start-maximized")
    options.add_argument(f"--user-data-dir={unique_user_data_dir}") 
    options.add_argument("--disable-blink-features=AutomationControlled")
    options.add_experimental_option("useAutomationExtension", False)
    options.add_experimental_option('excludeSwitches', ['enable-automation'])
    options.add_argument("--headless")
    options.add_argument("--no-sandbox")
    options.add_argument("--disable-dev-shm-usage")

    prefs = {
        "download.default_directory": user_download_dir,
        "download.prompt_for_download": False,
        "download.directory_upgrade": True,
        "safebrowsing.enabled": True,
        "profile.default_content_setting_values.automatic_downloads": 1,
        "profile.content_settings.exceptions.automatic_downloads.*.setting": 1,
        "plugins.always_open_pdf_externally": True,
        "credentials_enable_service": False,
        "profile.password_manager_enabled": False,
        "profile.block_third_party_cookies": False,
        "profile.default_content_settings.popups": 0,
        "profile.default_content_setting_values.notifications": 2
    }
    options.add_experimental_option("prefs", prefs)
    
    return options
def setup_driver(username=None, site_name="default_site"):
    user_download_dir = None
    if username:
        user_download_dir = setup_user_directory(username, site_name)
    else:
        user_download_dir = setup_user_directory()
    
    options = setup_chrome_options(user_download_dir)

    return webdriver.Chrome(options=options)

def login_user(driver, username, password):
    try:
        driver.get("https://eportal.incometax.gov.in/iec/foservices/#/login")
        
        # Enter Username
        panAdhaarUserId = WebDriverWait(driver,20).until(
            EC.visibility_of_element_located((By.ID, "panAdhaarUserId"))
        )
        panAdhaarUserId.clear()
        panAdhaarUserId.send_keys(username)
        panAdhaarUserId.send_keys(Keys.RETURN)

        # Handle Checkbox
        passwordCheckBox = WebDriverWait(driver,20).until(
            EC.visibility_of_element_located((By.ID, "passwordCheckBox"))
        )
        passwordCheckBox.click()

        # Handle Password
        loginPasswordField = WebDriverWait(driver, 20).until(
            EC.visibility_of_element_located((By.ID, "loginPasswordField"))
        )
        loginPasswordField.clear()
        loginPasswordField.send_keys(password)
        time.sleep(2)

        # Click On Continue
        continueBtn = WebDriverWait(driver, 20).until(EC.visibility_of_element_located((By.CSS_SELECTOR, "button.large-button-primary")))
        time.sleep(2)
        continueBtn.send_keys(Keys.RETURN)
        

        try:
            # Wait for the modal to appear and the "Login Here" button within it
            modal_login_button = WebDriverWait(driver, 20).until(
                EC.visibility_of_element_located(
                    (By.CSS_SELECTOR, "div.modal-footer button.defaultButton.primaryButton.primaryBtnMargin")
                )
            )
            modal_login_button.click()  # Click the button within the modal
        except Exception as e:
            print(f"Error: {e}")
            # Handle the case where neither element is found, if needed

        return True
    
    except Exception as e:
        print(f"Login failed for user {username}: {str(e)}")
        return False



def find_latest_download(directory):
    """Find the most recently downloaded file in the directory"""
    try:
        files = [os.path.join(directory, f) for f in os.listdir(directory) 
                if os.path.isfile(os.path.join(directory, f))]
        if not files:
            return None
        return max(files, key=os.path.getctime)
    except Exception as e:
        print(f"Error finding latest download: {str(e)}")
        return None

def check_response_to_outstanding_demand_page(driver, client_name):
    try:
        # Get the user directory for saving files
        client = frappe.get_doc("Client", client_name, fields=['username'])
        download_dir = setup_user_directory(client.username, "incometax")
        
        # Click Pending Actions
        pending_actions = WebDriverWait(driver, 10).until(
            EC.visibility_of_element_located((By.ID, "Pending Actions"))
        )
        pending_actions.click()
        time.sleep(2)
        
        # Click Worklist
        worklist_button = WebDriverWait(driver, 10).until(
            EC.visibility_of_element_located((By.XPATH, "//span[contains(text(), 'Response to Outstanding Demand')]"))
        )
        worklist_button.click()
        time.sleep(3)
        
        # Check page content
        page_source = driver.page_source
        soup = BeautifulSoup(page_source, 'html.parser')
        
        if "No records of outstanding demand found" in soup.get_text():
            return "No Action Required", []
        
        demands = extract_demand_details(soup)
        if not demands:
            return "No Demands Found", []
        
        # Get download buttons
        download_buttons = driver.find_elements(By.XPATH, 
            "//button[contains(@class, 'downloadiconPosition') and contains(@class, 'secondaryButton')]"
        )
        
        if not download_buttons:
            return "No Download Buttons Found", demands
        
        # Process each download
        for i, button in enumerate(download_buttons):
            if i >= len(demands):
                break
                
            # Get latest file before download
            before_download = find_latest_download(download_dir)
            before_time = time.time()
            
            # Click download button
            driver.execute_script("arguments[0].scrollIntoView({ behavior: 'smooth', block: 'center' });", button)
            time.sleep(2)
                    
                    # Click download
            download_button = WebDriverWait(driver, 10).until(
                        EC.element_to_be_clickable(button)
                    )
            download_button.click()
            
            print(download_dir)
            # Wait for and get new download
            file_path = wait_for_download_using_latest(download_dir, before_download, before_time)
            print(file_path)
            if file_path:
                # Successfully downloaded
                create_or_update_demand(demands[i], client_name, file_path)
                demands[i]['download_status'] = 'Downloaded and Attached'
            else:
                # Download failed
                create_or_update_demand(demands[i], client_name, None)
                demands[i]['download_status'] = 'Download Failed'
                
        return "Action Required", demands
            
    except Exception as e:
        print(f"Error in check_response_to_outstanding_demand_page: {str(e)}")
        return f"Error: {str(e)}", []

def wait_for_download_using_latest(download_dir, before_download, before_time, timeout=15):
    """Wait for download using find_latest_download"""
    start_time = time.time()
    
    while time.time() - start_time < timeout:
        # Find most recent file
        latest_file = find_latest_download(download_dir)
        
        # Check if we have a new file
        if latest_file and (not before_download or 
                           os.path.basename(latest_file) != os.path.basename(before_download) or
                           os.path.getctime(latest_file) > before_time):
            
            # Wait a moment for download to complete
            time.sleep(1)
            
            # Verify file is valid
            if os.path.exists(latest_file) and os.path.getsize(latest_file) > 0:
                print(f"New file downloaded: {os.path.basename(latest_file)}")
                return latest_file
                
        time.sleep(0.5)
        
    print("No new file detected after download attempt")
    return None

# def setup_user_directory(username):
#     """Create and return user download directory"""
#     base_dir = os.path.join(frappe.utils.get_site_path(), "downloads")
#     user_dir = os.path.join(base_dir, username)
#     os.makedirs(user_dir, exist_ok=True)
#     return user_dir




def create_or_update_demand_old(demand, client_name, file_path):
    """Create or update demand document and attach downloaded notice"""
    try:
        if not demand:
            print("Warning: Empty demand data provided")
            return
            
        # Clean the amount string to get just the number
        amount_str = demand.get('outstanding_demand_amount', '0')
        amount = int(re.sub(r"[^\d]", "", amount_str))
        
        # Check for existing document
        response_docs = frappe.get_all("Response to Outstanding Demand", 
            filters={
                'demand_reference_no': demand.get('demand_reference_no', ''),
                'assessment_year': demand.get('assessment_year', ''),
                'client': client_name
            }
        )
        
        is_new_doc = False
        if response_docs:
            # Update existing document
            doc = frappe.get_doc("Response to Outstanding Demand", response_docs[0].name)
        else:
            # Create new document
            is_new_doc = True
            doc = frappe.new_doc("Response to Outstanding Demand")
            doc.demand_reference_no = demand.get('demand_reference_no', '')
            doc.assessment_year = demand.get('assessment_year', '')
            doc.outstanding_demand_amount = amount
            doc.section_code = demand.get('section_code', '')
            doc.rectification_rights = demand.get('rectification_rights', '')
            doc.mode_of_service = demand.get('mode_of_service', '')
            doc.response_type = demand.get('response_type', '')
            doc.client = client_name
            
            # Save the document first to get a name
            doc.insert()
        
        if file_path and os.path.exists(file_path):
            try:
                # For new docs, we've already saved above
                # For existing docs, handle attachments first
                if not is_new_doc:
                    # Get existing attachments
                    existing_attachments = frappe.get_all("File", 
                        filters={
                            "attached_to_doctype": "Response to Outstanding Demand",
                            "attached_to_name": doc.name,
                            "attached_to_field": "notice_letter"
                        }
                    )
                    
                    # Remove existing attachment if any
                    for attachment in existing_attachments:
                        frappe.delete_doc('File', attachment.name)
                
                # Add new attachment
                filename = os.path.basename(file_path)
                
                # Read file content
                file_content = None
                with open(file_path, 'rb') as f:
                    file_content = f.read()
                
                if file_content:
                    # Create file attachment
                    attachment = frappe.get_doc({
                        'doctype': 'File',
                        'file_name': filename,
                        'content': file_content,
                        'attached_to_doctype': 'Response to Outstanding Demand',
                        'attached_to_name': doc.name,
                        'attached_to_field': 'notice_letter'
                    })
                    attachment.insert()
                    
                    # Update document with file reference
                    doc.notice = attachment.file_url
                    doc.save()
                    
                    print(f"Successfully attached {filename} to document {doc.name}")
                else:
                    print(f"Warning: Empty file content for {file_path}")
            except Exception as e:
                print(f"Error attaching file: {str(e)}")
                # Continue even if attachment fails
        elif file_path:
            print(f"Warning: File path provided but file doesn't exist: {file_path}")
        
        # Final save and commit
        if not is_new_doc:
            doc.save()
        frappe.db.commit()
        
        return doc.name
        
    except Exception as e:
        error_msg = f"Error processing demand document: {str(e)}"
        print(error_msg)
        frappe.log_error(error_msg)
        return None




def process_pdf_file(file_path, client_name):
    """
    Process PDF file and handle encryption if necessary.
    Returns file content and encryption status.
    """
    try:
        with open(file_path, 'rb') as f:
            reader = PyPDF2.PdfReader(f)
            
            if not reader.is_encrypted:
                return f.read(), False, "File processed successfully"
                
            # Handle encrypted PDF
            try:
                client_doc = frappe.get_doc("Client", client_name)
                password = f"{client_doc.username.lower()}{client_doc.dob}"
                reader.decrypt(password)
                
                writer = PyPDF2.PdfWriter()
                for page in reader.pages:
                    writer.add_page(page)
                
                output_buffer = BytesIO()
                writer.write(output_buffer)
                return output_buffer.getvalue(), True, "File decrypted successfully"
                
            except (PdfReadError, ValueError) as e:
                f.seek(0)
                return f.read(), False, f"Unable to decrypt file: {str(e)}"
                
    except Exception as e:
        raise Exception(f"Error processing PDF: {str(e)}")

def handle_file_attachment(doc, file_path, file_content):
    """
    Handle file attachment for the document.
    Manages existing attachments and creates new ones.
    """
    try:
        # Remove existing attachments
        existing_attachments = frappe.get_all("File", 
            filters={
                "attached_to_doctype": "Response to Outstanding Demand",
                "attached_to_name": doc.name,
                "attached_to_field": "notice_letter"
            }
        )
        
        for attachment in existing_attachments:
            frappe.delete_doc('File', attachment.name)
        
        # Create new attachment
        filename = os.path.basename(file_path)
        attachment = frappe.get_doc({
            'doctype': 'File',
            'file_name': filename,
            'content': file_content,
            'attached_to_doctype': 'Response to Outstanding Demand',
            'attached_to_name': doc.name,
            'attached_to_field': 'notice',
            'is_private': 1
        })
        attachment.insert()
        
        return attachment.file_url
        
    except Exception as e:
        raise Exception(f"Error handling attachment: {str(e)}")

def create_demand_doc(demand, client_name):
    """
    Create new demand document with given data.
    """
    try:
        amount_str = demand.get('outstanding_demand_amount', '0')
        amount = int(re.sub(r"[^\d]", "", amount_str))
        
        doc = frappe.new_doc("Response to Outstanding Demand")
        doc.demand_reference_no = demand.get('demand_reference_no', '')
        doc.assessment_year = demand.get('assessment_year', '')
        doc.outstanding_demand_amount = amount
        doc.section_code = demand.get('section_code', '')
        doc.rectification_rights = demand.get('rectification_rights', '')
        doc.mode_of_service = demand.get('mode_of_service', '')
        doc.response_type = demand.get('response_type', '')
        doc.client = client_name
        
        doc.insert()
        return doc
        
    except Exception as e:
        raise Exception(f"Error creating demand document: {str(e)}")

def create_or_update_demand(demand, client_name, file_path):
    """
    Main function to create or update demand document and handle attachments.
    """
    try:
        if not demand:
            print("Warning: Empty demand data provided")
            return

        # Check for existing document
        response_docs = frappe.get_all("Response to Outstanding Demand", 
            filters={
                'demand_reference_no': demand.get('demand_reference_no', ''),
                'assessment_year': demand.get('assessment_year', ''),
                'client': client_name
            }
        )
        
        # Get or create document
        if response_docs:
            doc = frappe.get_doc("Response to Outstanding Demand", response_docs[0].name)
        else:
            doc = create_demand_doc(demand, client_name)
        
        # Handle file attachment if provided
        if file_path and os.path.exists(file_path):
            try:
                # Process PDF file
                file_content, was_decrypted, status = process_pdf_file(file_path, client_name)
                print(status)
                
                if file_content:
                    # Attach file and update document
                    file_url = handle_file_attachment(doc, file_path, file_content)
                    doc.notice = file_url
                    doc.save()
                    
                    status_msg = f"File attached successfully"
                    if was_decrypted:
                        status_msg += " (decrypted)"
                    print(status_msg)
                    
            except Exception as e:
                print(f"Error handling file: {str(e)}")
                frappe.log_error(str(e), "Demand File Processing")
        
        frappe.db.commit()
        return doc.name
        
    except Exception as e:
        error_msg = f"Error in demand processing: {str(e)}"
        print(error_msg)
        frappe.log_error(error_msg)
        return None


def extract_demand_details(soup):
    """Extract demand details from the page"""
    demands = []
    
    try:
        demand_cards = soup.select("div.card-container.matCard.mobWid.mt-4")
        
        for card in demand_cards:
            demand = {}
            
            # Extract Demand Reference No
            ref_span = card.select_one("span.heading5.mNoWrap.m-sm-wordWrap")
            if ref_span:
                demand["demand_reference_no"] = ref_span.text.strip()
            
            # Extract Assessment Year
            ass_yr_span = card.select_one("div.ass_yr_spacing span.heading5")
            if ass_yr_span:
                demand["assessment_year"] = ass_yr_span.text.strip()

            # Extract Outstanding Amount
            amount_div = card.select_one("div.heading6")
            if amount_div:
                demand["outstanding_demand_amount"] = amount_div.text.strip()

            # Extract other details
            details_div = card.select_one("div.col-lg-3.pipeline.row-part3-large")
            if details_div:
                detail_rows = details_div.select("div.pb12")
                
                for row in detail_rows:
                    label = row.select_one("span.dataHeading")
                    value = row.select_one("span.subtitle2") or row.select_one("span.body2")
                    
                    if label and value:
                        key = label.text.split(":")[0].strip().lower().replace(" ", "_")
                        val = value.text.strip()
                        demand[key] = val

            if demand.get("demand_reference_no"):
                demands.append(demand)
    
    except Exception as e:
        print(f"Error in demand extraction: {str(e)}")
    
    return demands


def logout_user(driver):
    try:
        print("logout .......")
        driver.execute_script("document.body.scrollTop = 0; document.documentElement.scrollTop = 0;")

        # Allow time for scrolling
        time.sleep(1)

        profileMenubtn = WebDriverWait(driver, 20).until(
            EC.element_to_be_clickable((By.CSS_SELECTOR, "button.profileMenubtn"))
        )
        profileMenubtn.click()

        logOut = WebDriverWait(driver, 20).until(
            EC.element_to_be_clickable((By.XPATH, "//span[contains(text(), 'Log Out')]"))
        )
        logOut.click()
        time.sleep(3)
        return True
    except Exception as e:
        print(f"Logout failed: {str(e)}")
        return False


def setup_user_directory(username="user_not_defined", site_name="default_site"):
    """Create and return a user-specific download directory under a site-specific path"""
    
    # Create base downloads directory if it doesn't exist
    base_download_dir = os.path.join(os.getcwd(), "downloads")
    if not os.path.exists(base_download_dir):
        os.makedirs(base_download_dir)
    
    # Create site-specific directory
    site_dir = os.path.join(base_download_dir, site_name)
    if not os.path.exists(site_dir):
        os.makedirs(site_dir)
    
    # Create user-specific directory under the site directory
    user_dir = os.path.join(site_dir, username)
    if not os.path.exists(user_dir):
        os.makedirs(user_dir)
    
    return user_dir



def handle_eproceedings_downloads(driver, username):

    def click_back_button(driver, max_attempts=3):
        """Helper function to click back button with retries"""
        for attempt in range(max_attempts):
            try:
                back_button = WebDriverWait(driver, 20).until(
                    EC.visibility_of_element_located((By.XPATH, "//button[contains(@class, 'large-button-secondary') and contains(@class, 'previousIcon') and contains(text(), 'Back')]"))
                )
                driver.execute_script("arguments[0].scrollIntoView({block: 'center'});", back_button)
                time.sleep(2)
                try:
                    back_button.click()
                except:
                    driver.execute_script("arguments[0].click();", back_button)
                time.sleep(3)
                return True
            except Exception as e:
                print(f"Back button click attempt {attempt + 1} failed: {str(e)}")
                time.sleep(2)
        return False

    def extract_container_details(container_element):
        """Extract all details from a single notice container."""
        print("container")
        print(container_element)
        try:
            details = {}
            
            # Extract notice reference ID
            ref_id_elem = container_element.select_one("span:-soup-contains('Notice/ Communication Reference ID :')")
            details['notice_din'] = None

            if ref_id_elem:
                ref_id = ref_id_elem.find_next('span', class_='heading5')
                if ref_id:
                    details['notice_din'] = ref_id.text.strip()

            if not details['notice_din']:
                ref_id_elem = container_element.select_one("span:-soup-contains(' Document Identification Number (DIN) : ')")
                if ref_id_elem:
                    doc_id = ref_id_elem.find_next('span', class_='heading5')
                    if doc_id:
                        details['notice_din'] = doc_id.text.strip()

            
            # Extract notice number and type
            notice_elem = container_element.select_one("div.heading6")
            if notice_elem and notice_elem.text.strip().isdigit():
                details['notice_section'] = notice_elem.text.strip()
                details['notice_number'] = notice_elem.text.strip()
                

            # Step 1: First find the div that contains the "Document reference ID" text
            doc_ref_container = container_element.find("div", class_="dataHeading mt-1", text=lambda t: "Document reference ID" in t if t else False)

            if doc_ref_container:
                # Step 2: Get the parent div
                parent_div = doc_ref_container.parent
                
                if parent_div:
                    # Step 3: From the parent div, find the heading6 div
                    doc_id_elem = parent_div.find("div", class_="heading6 lineSpace")
                    
                    if doc_id_elem:
                        doc_id = doc_id_elem.get_text(strip=True)
                        details['notice_communication_reference_id'] = doc_id.strip()
                        print("✅ Extracted Document Reference ID:", doc_id)
                    else:
                        print("❌ No <div class='heading6 lineSpace'> found inside parent")
                else:
                    print("❌ Couldn't find parent of Document reference ID div")
            else:
                print("❌ Couldn't find div containing 'Document reference ID' text")
            
            
            # Extract description
            desc_elem = container_element.select_one("span:-soup-contains('Description :')")
            if desc_elem:
                description = desc_elem.find_next('span', class_='subtitle2')
                if description:
                    details['description'] = description.text.strip()
            
            # Extract dates
            issued_date_elem = container_element.select_one("span:-soup-contains('Issued On :')")
            if issued_date_elem:
                issued_date = issued_date_elem.find_next('span', class_='subtitle2')
                if issued_date:
                    notice_sent_date_value = issued_date.text.strip()
                    converted_date = MakeDateValid.get_frappe_date(notice_sent_date_value)
                    details['notice_sent_date'] = converted_date

                    
            due_date_elem = container_element.select_one("span:-soup-contains('Response Due Date :')")
            if due_date_elem:
                due_date = due_date_elem.find_next('span', class_='subtitle2')
                if due_date:
                    response_due_date_value = due_date.text.strip()
                    converted_date = MakeDateValid.get_frappe_date(response_due_date_value)
                    details['response_due_date'] = converted_date

            
            print("Container Details:")
            print(details)
            print()
            
            return details
        except Exception as e:
            print(f"Error extracting container details: {str(e)}")
            return {}

    def extract_header_details(driver):
        """Extract header details common to all notices on the page"""
        try:
            soup = BeautifulSoup(driver.page_source, 'html.parser')
            header_details = {}
            
            # Extract Proceeding Name
            proc_name_elem = soup.select_one("div.dataHeading:-soup-contains('Proceeding Name')")
            if proc_name_elem:
                value = proc_name_elem.find_next('div', class_='subtitle1')
                if value:
                    header_details['proceeding_name'] = value.text.strip()

            # Extract Assessment Year
            assessment_elem = soup.select_one("div.dataHeading:-soup-contains('Assessment Year')")
            if assessment_elem:
                value = assessment_elem.find_next('div', class_='subtitle1')
                if value:
                    header_details['assessment_year'] = value.text.strip()
            
            # Extract Financial Year
            financial_elem = soup.select_one("div.dataHeading:-soup-contains('Financial Year')")
            if financial_elem:
                value = financial_elem.find_next('div', class_='subtitle1')
                if value:
                    header_details['financial_year'] = value.text.strip()
            
            # Add client details
            # client = frappe.get_doc("Client", username, fields=['username'])
            header_details['client'] = username
            header_details['proceeding_status'] = 'Active'

            print("HEADER Details:")
            print(header_details)
            print()
                    
            return header_details
        except Exception as e:
            print(f"Error extracting header details: {str(e)}")
            return {}


    def process_page2_buttons(driver, button_type="response"):
        """Helper function to process buttons on page 2 and create/update E Proceeding records"""
        try:
            # Extract header details first (common to all notices on this page)
            header_details = extract_header_details(driver)
            
            if button_type == "response":
                # Find View Response buttons
                buttons = WebDriverWait(driver, 20).until(
                    EC.presence_of_all_elements_located((By.XPATH, "//button[contains(@class, 'defaultButton primaryButton cardbutton') and contains(text(), 'View Response')]"))
                )
                button_name = "View Response"
            else:
                # Find Notice/Letter buttons
                buttons = WebDriverWait(driver, 20).until(
                    EC.presence_of_all_elements_located((By.XPATH, "//button[contains(@class, 'defaultButton secondaryButton mMt16 cardbutton') and contains(text(), 'Notice/Letter pdf')]"))
                )
                button_name = "Notice/Letter"

            if not buttons:
                print(f"No {button_name} buttons found")
                return 0, []

            print(f"Found {len(buttons)} {button_name} buttons")
            processed = 0
            created_or_updated_docs = []

            for j in range(len(buttons)):
                try:
                    # Get page source for extraction before clicking button
                    soup = BeautifulSoup(driver.page_source, 'html.parser')
                    
                    # Re-find buttons to avoid stale elements
                    current_buttons = WebDriverWait(driver, 20).until(
                        EC.presence_of_all_elements_located((By.XPATH, 
                            "//button[contains(@class, 'defaultButton primaryButton cardbutton') and contains(text(), 'View Response')]" if button_type == "response"
                            else "//button[contains(@class, 'defaultButton secondaryButton mMt16 cardbutton') and contains(text(), 'Notice/Letter pdf')]"
                        ))
                    )
                    if j >= len(current_buttons):
                        continue
                    
                    button = current_buttons[j]
                    
                    # Find the container for this specific button
                    card_container = button.find_element(By.XPATH, "./ancestor::div[contains(@class, 'card-container')]")
                    card_html = card_container.get_attribute('outerHTML')
                    container_soup = BeautifulSoup(card_html, 'html.parser')
                    
                    # Extract container details
                    container_details = extract_container_details(container_soup)
                    
                    # Merge header and container details
                    proceeding_details = {**header_details, **container_details}
                    
                    print(f"Processing {button_name} button {j + 1}/{len(buttons)}")
                    
                    # Click button to go to third page
                    driver.execute_script("arguments[0].scrollIntoView({block: 'center'});", button)
                    time.sleep(2)
                    driver.execute_script("arguments[0].click();", button)
                    time.sleep(3)
                    
                    # Determine file type and download
                    file_path = None
                    if button_type == "response":
                        # Find and click all hyperLink elements
                        try:
                            hyperlink_elements = WebDriverWait(driver, 10).until(
                                EC.presence_of_all_elements_located((By.XPATH, 
                                    "//span[contains(@class, 'subtitle1')]/span[contains(@class, 'hyperLink') and @role='link']"))
                            )
                            
                            print(f"Found {len(hyperlink_elements)} hyperLink elements")
                            
                            # Get the user directory for saving files
                            client = frappe.get_doc("Client", username, fields=['username'])
                            user_dir = setup_user_directory(client.username, "incometax")
                            
                            for i, link in enumerate(hyperlink_elements):
                                try:
                                    link_text = link.text
                                    print(f"Clicking hyperLink {i+1}/{len(hyperlink_elements)}: {link_text}")
                                    driver.execute_script("arguments[0].scrollIntoView({block: 'center'});", link)
                                    time.sleep(1)
                                    driver.execute_script("arguments[0].click();", link)
                                    time.sleep(7)
                                    # Wait for download to complete                                    
                                    
                                    # Find the latest downloaded file
                                    file_path = find_latest_download(user_dir)
                                    print("file_path", file_path)
                                    print("user_Dir", user_dir)
                                    if file_path:
                                        response = add_hyperlink_to_eproceeding_replies(proceeding_details, link_text, file_path)
                                        print(response)
                                    else:
                                        print(f"No file downloaded for hyperLink {i+1}")
                                        
                                except Exception as e:
                                    print(f"Error clicking hyperLink {i+1}: {str(e)}")
                        except Exception as e:
                            print(f"Error finding hyperLink elements: {str(e)}")
                        # END NEW CODE


                        # Handle Acknowledgement button
                        try:
                            download_button = WebDriverWait(driver, 20).until(
                                EC.visibility_of_element_located((By.XPATH, "//button[contains(@class, 'normal-button-secondary iconBefore downloadIcon') and contains(text(), 'Acknowledgement')]"))
                            )
                            print("Clicking Acknowledgement button")
                            
                            # Get the user directory for saving files
                            client = frappe.get_doc("Client", username, fields=['username'])
                            user_dir = setup_user_directory(client.username, "incometax")
                            
                            # Download the file
                            driver.execute_script("arguments[0].scrollIntoView({block: 'center'});", download_button)
                            time.sleep(2)
                            driver.execute_script("arguments[0].click();", download_button)
                            time.sleep(5)  # Wait for download to complete
                            
                            # Find the latest downloaded file
                            file_path = find_latest_download(user_dir)
                            file_type = 'response_acknowledgement'
                            
                        except Exception as e:
                            print(f"Error finding Acknowledgement button: {str(e)}")
                            if not click_back_button(driver):
                                raise Exception("Failed to return to page 2")
                            continue
                    else:
                        # Handle Download button for Notice/Letter
                        try:
                            download_button = WebDriverWait(driver, 20).until(
                                EC.visibility_of_element_located((By.XPATH, "//button[contains(@class, 'normal-button-secondary iconBefore downloadPdf downloadIcon') and contains(text(), 'Download')]"))
                            )
                            print("Clicking Download button")
                            
                            # Get the user directory for saving files
                            client = frappe.get_doc("Client", username, fields=['username'])
                            user_dir = setup_user_directory(client.username, "incometax")
                            
                            # Download the file
                            driver.execute_script("arguments[0].scrollIntoView({block: 'center'});", download_button)
                            time.sleep(2)
                            driver.execute_script("arguments[0].click();", download_button)
                            time.sleep(5)  # Wait for download to complete
                            
                            # Find the latest downloaded file
                            file_path = find_latest_download(user_dir)
                            file_type = 'notice_letter'
                            
                        except Exception as e:
                            print(f"Error finding Download button: {str(e)}")
                            if not click_back_button(driver):
                                raise Exception("Failed to return to page 2")
                            continue
                    
                    # Create or update E Proceeding record and attach the file
                    # if proceeding_details.get('notice_communication_reference_id') or if proceeding_details.get('notice_din') :
                    if True:
                        doc_name = create_or_update_eproceeding(proceeding_details, file_path, file_type)
                        if doc_name:
                            created_or_updated_docs.append(doc_name)
                    
                    processed += 1
                    
                    # Return to page 2
                    if not click_back_button(driver):
                        raise Exception(f"Failed to return to page 2 after {button_name}")
                    time.sleep(2)
                    
                except Exception as e:
                    print(f"Error processing {button_name} button {j + 1}: {str(e)}")
                    continue

            return processed, created_or_updated_docs

        except Exception as e:
            print(f"Error in processing {button_type} buttons: {str(e)}")
            return 0, []

    def find_latest_download(directory):
        """Find the most recently downloaded file in the directory"""
        try:
            files = [os.path.join(directory, f) for f in os.listdir(directory) if os.path.isfile(os.path.join(directory, f))]
            if not files:
                return None
            return max(files, key=os.path.getctime)
        except Exception as e:
            print(f"Error finding latest download: {str(e)}")
            return None


    def create_or_update_eproceeding(details, file_path=None, file_type=None):
        """Create or update an E Proceeding record with retry logic"""
        max_retries = 1
        retry_count = 0
        
        while retry_count < max_retries:
            try:
                # Begin a new transaction
                # frappe.db.begin()
                
                # Check if record exists
                filters = {}

                # Check if 'notice_communication_reference_id' is available
                if details.get('notice_din'):
                    filters["notice_din"] = details['notice_din']
                elif details.get('notice_communication_reference_id'):
                    filters["notice_communication_reference_id"] = details['notice_communication_reference_id']
                else:
                    # If not available, use other fields for lookup
                    filters.update({
                        "assessment_year": details.get('assessment_year', ""),
                        "proceeding_name": details.get('proceeding_name', ""),
                        "notice_section": details.get('notice_section', ""),
                        "notice_sent_date": details.get('notice_sent_date', ""),
                        "response_due_date": details.get('response_due_date', ""),
                        "client": details.get('client', ""),
                    })

                    # Remove empty fields to avoid unnecessary filters
                    filters = {k: v for k, v in filters.items() if v}

                # filters["client"] = details['client']
                # Fetch records based on available filters
                existing_records = frappe.get_all(
                    "E Proceeding", 
                    filters=filters,
                    fields=["name"]
                )

                # proceeding_name notice_section
                if existing_records:
                    doc = frappe.get_doc("E Proceeding", existing_records[0]['name'])
                    
                    # Update fields
                    for key, value in details.items():
                        if hasattr(doc, key) and value:  # Only update if value is provided
                            setattr(doc, key, value)
                    
                    # Attach file if provided
                    if file_path and file_type:
                        success = attach_file_to_doc(doc.doctype, doc.name, file_path, file_type)
                        if not success:
                            frappe.db.rollback()
                            raise Exception("File attachment failed")
                    
                    doc.save()
                    frappe.db.commit()
                    print(f"Updated E Proceeding record: {doc.name}")
                    
                else:
                    # Create new record
                    doc = frappe.get_doc({
                        "doctype": "E Proceeding",
                        **details
                    })
                    doc.insert()
                    frappe.db.commit()
                    
                    # Attach file if provided
                    if file_path and file_type:
                        success = attach_file_to_doc(doc.doctype, doc.name, file_path, file_type)
                        if not success:
                            frappe.db.rollback()
                            raise Exception("File attachment failed")
                    
                    doc.save()
                    frappe.db.commit()
                    print(f"Created new E Proceeding record: {doc.name}")

                return doc.name
                
            except Exception as e:
                frappe.db.rollback()
                print(f"Attempt {retry_count + 1} failed: {str(e)}")
                retry_count += 1
                time.sleep(2)  # Wait before retry
                
                if retry_count == max_retries:
                    print(f"All attempts failed for notice ID: {details.get('notice_communication_reference_id')}")
                    return None

    def attach_file_to_doc(doctype, docname, file_path, field_name):
        """Attach file to document with improved error handling"""
        try:
            if not file_path or not os.path.exists(file_path):
                print(f"File not found: {file_path}")
                return False
                
            file_name = os.path.basename(file_path)
            
            # Read file content
            with open(file_path, "rb") as f:
                content = f.read()
            
            # Create File document
            file_doc = frappe.get_doc({
                "doctype": "File",
                "file_name": file_name,
                "attached_to_doctype": doctype,
                "attached_to_name": docname,
                "attached_to_field": field_name,
                "is_private": 1,
                "content": content
            })
            
            file_doc.insert()
            frappe.db.commit()
            
            # Update the main document
            doc = frappe.get_doc(doctype, docname)
            setattr(doc, field_name, file_doc.file_url)
            doc.save()
            frappe.db.commit()
            
            print(f"Successfully attached file {file_name} to {doctype} {docname}")
            return True
            
        except Exception as e:
            print(f"Error attaching file: {str(e)}")
            return False


    try:
        processed_count = 0
        all_created_or_updated_docs = []
        
        # Find all "View Notices/Orders" buttons on page 1
        try:
            view_notices_buttons = WebDriverWait(driver, 20).until(
                EC.presence_of_all_elements_located((By.XPATH, "//button[contains(@class, 'defaultButton primaryButton') and contains(text(), 'View Notices/Orders')]"))
            )

            if not view_notices_buttons:
                print("No View Notices/Orders buttons found")
                return "No items to process"

            print(f"\nFound {len(view_notices_buttons)} items on page 1")

            # Process each View Notices/Orders button
            for i in range(len(view_notices_buttons)):
                try:
                    # Re-find buttons to avoid stale elements
                    current_buttons = WebDriverWait(driver, 20).until(
                        EC.presence_of_all_elements_located((By.XPATH, "//button[contains(@class, 'defaultButton primaryButton') and contains(text(), 'View Notices/Orders')]"))
                    )
                    if i >= len(current_buttons):
                        continue
                        
                    notice_button = current_buttons[i]
                    
                    # Click View Notices/Orders to go to page 2
                    driver.execute_script("arguments[0].scrollIntoView({block: 'center'});", notice_button)
                    time.sleep(2)
                    driver.execute_script("arguments[0].click();", notice_button)
                    time.sleep(3)

                    # Process all View Response buttons first
                    response_count, response_docs = process_page2_buttons(driver, "response")
                    processed_count += response_count
                    all_created_or_updated_docs.extend(response_docs)
                    
                    # Then process all Notice/Letter buttons
                    notice_count, notice_docs = process_page2_buttons(driver, "notice")
                    processed_count += notice_count
                    all_created_or_updated_docs.extend(notice_docs)

                    # Return to page 1
                    if not click_back_button(driver):
                        raise Exception("Failed to return to page 1")
                    time.sleep(2)

                except Exception as e:
                    print(f"Error processing View Notices/Orders button {i + 1}: {str(e)}")
                    continue

        except Exception as e:
            print(f"Error finding View Notices/Orders buttons: {str(e)}")
            return f"Error: {str(e)}"

        # Return summary of processing
        return f"Downloads Completed - Processed {processed_count} items, Created/Updated {len(set(all_created_or_updated_docs))} E Proceeding records"
        
    except Exception as e:
        print(f"Error in handle_eproceedings_downloads: {str(e)}")
        return f"Error: {str(e)}"



def add_hyperlink_to_eproceeding_replies(proceeding_details, link_text, file_path):
    """
    Add a hyperlink file to the replies child table of an E Proceeding.
    Creates a new E Proceeding if none exists matching the given details.
    
    Args:
        proceeding_details (dict): Details to identify/create the E Proceeding document
        link_text (str): Text of the hyperlink (to be used as file_name)
        file_path (str): Path to the downloaded file
        
    Returns:
        tuple: (bool, str) - (Success status, Message describing the outcome)
    """
    try:
        # Validate input parameters
        if not all([proceeding_details, link_text, file_path]):
            return False, "Missing required parameters"
            
        if not os.path.exists(file_path):
            return False, f"File not found: {file_path}"
            
        # Prepare filters for finding existing E Proceeding
        filters = _prepare_filters(proceeding_details)
        
        # Get or create E Proceeding document
        doc = _get_or_create_eproceeding(filters, proceeding_details)
        if not doc:
            return False, "Failed to get or create E Proceeding document"
            
        # Check for existing file
        if _file_exists_in_replies(doc, link_text):
            return True, f"File '{link_text}' already exists in E Proceeding {doc.name}"
            
        # Create and attach new file
        success, message = _create_and_attach_file(doc, link_text, file_path)
        return success, message
        
    except Exception as e:
        frappe.db.rollback()
        return False, f"Error processing request: {str(e)}"

def _prepare_filters(proceeding_details):
    """Prepare filters for finding existing E Proceeding."""
    if proceeding_details.get('notice_din'):
        return {"notice_din": proceeding_details['notice_din']}
    
    if proceeding_details.get('notice_communication_reference_id'):
        return {"notice_communication_reference_id": proceeding_details['notice_communication_reference_id']}
    
    # Use other fields for lookup
    filter_fields = [
        "assessment_year",
        "proceeding_name",
        "notice_section",
        "notice_sent_date",
        "response_due_date",
        "client"
    ]
    
    return {
        k: proceeding_details.get(k, "")
        for k in filter_fields
        if proceeding_details.get(k)
    }

def _get_or_create_eproceeding(filters, proceeding_details):
    """Get existing E Proceeding or create new one."""
    existing_records = frappe.get_all(
        "E Proceeding",
        filters=filters,
        fields=["name"]
    )
    
    if existing_records:
        return frappe.get_doc("E Proceeding", existing_records[0]['name'])
        
    # Create new record
    try:
        new_doc = frappe.get_doc({
            "doctype": "E Proceeding",
            **proceeding_details
        })
        new_doc.insert()
        frappe.db.commit()
        return new_doc
    except Exception as e:
        frappe.log_error(f"Failed to create E Proceeding: {str(e)}")
        return None

def _file_exists_in_replies(doc, link_text):
    """Check if file already exists in replies."""
    return any(reply.file_name == link_text for reply in doc.replies)

def _create_and_attach_file(doc, link_text, file_path):
    """Create File document and attach to E Proceeding."""
    try:
        # First, create a new E Proceeding Reply
        reply_doc = frappe.get_doc({
            "doctype": "E Proceeding Reply",
            "parent": doc.name,
            "parenttype": "E Proceeding",
            "parentfield": "replies",
            "file_name": link_text
        })
        reply_doc.insert()
        
        # Read file content
        with open(file_path, "rb") as f:
            content = f.read()
        
        # Create File document
        file_doc = frappe.get_doc({
            "doctype": "File",
            "file_name": link_text,
            "attached_to_doctype": "E Proceeding Reply",
            "attached_to_field": "file",
            "attached_to_name": reply_doc.name,  # Added this line
            "is_private": 1,
            "content": content
        })
        
        file_doc.insert()
        
        # Update the reply with the file URL
        reply_doc.file = file_doc.file_url
        reply_doc.save()
        
        frappe.db.commit()
        
        return True, f"Successfully added file '{link_text}' to E Proceeding {doc.name}"
        
    except Exception as e:
        frappe.db.rollback()
        return False, f"Failed to create and attach file: {str(e)}"


# Update the check_eproceedings_status function to use the new handler
def check_eproceedings_status(driver, client_name):
    try:
        driver.execute_script("document.body.scrollTop = 0; document.documentElement.scrollTop = 0;")

        # Allow time for scrolling
        time.sleep(1)
        # Click Pending Actions
        WebDriverWait(driver, 20).until(
            EC.visibility_of_element_located((By.ID, "Pending Actions"))
        ).click()

        # Click e-Proceedings
        e_proceedings_button = WebDriverWait(driver, 20).until(
            EC.visibility_of_element_located((By.XPATH, "//span[contains(text(), 'e-Proceedings')]"))
        )
        e_proceedings_button.click()
        
        # Wait for page to load
        time.sleep(3)
        
        page_source = driver.page_source
        soup = BeautifulSoup(page_source, 'html.parser')
        
        if "No e-Proceedings records are found" in soup.get_text():
            return "No Action Required"
        else:
            try:
                # Wait for and click the Excel download button
                download_button = WebDriverWait(driver, 20).until(
                    EC.element_to_be_clickable((By.CSS_SELECTOR, "button.secondaryButton.defaultButton.iconBefore.downloadSecNote.downloadButtonsec"))
                )
                download_button.click()
                
                # Wait for download to complete (adjust time if needed)
                time.sleep(2)
            except Exception as e:
                print(f"Failed to download Excel for user {client_name}: {str(e)}")
                # return "Action Required - Download Failed"

            try:    
                handle_eproceedings_downloads(driver, client_name)
            except Exception as e:
                print(f"Failed to handle the dowloads in eproceeding Excel for user {client_name}: {str(e)}")

            return "Action Required"
    except Exception as e:
        return f"Error: {str(e)}"



@frappe.whitelist()
def login_into_portal(client_name):
    """
    Queue background jobs for processing a selected client.
    
    Args:
        client_name: Client document name to process.
    """
    if client_name:
        # Get the client doc
        doc = frappe.get_doc("Client", client_name)

        # Queue the background job
        if not doc.disabled:
            frappe.enqueue(
                'fin_buddy.events.incometax_gov.login',
                queue='long',
                timeout=3000,
                client_name=client_name
            )
        else:
            return {"message": f"Client {client_name} is disabled !!"}
        
        return {"message": f"Successfully queued client {client_name} for processing"}
    else:
        return {"message": "No client found to process"}

def login(client_name):
    try:
        doc = frappe.get_doc("Client", client_name)
        # Get credentials
        username = doc.username
        password = doc.get_password("password")
        
        if not username or not password:
            frappe.log_error(
                f"Missing credentials for client {client_name}",
                "Income Tax Portal Sync Error"
            )
            return
        
        # Setup Selenium and process
        driver = setup_driver(username, "incometax")
                
        try:
            login_user(driver, username, password)
        except Exception as e:
            frappe.log_error("Login User Failed", str(e))
    
    except Exception as e:
        frappe.log_error("Login Error In Client", str(e))

@frappe.whitelist()
def process_selected_clients(client_names):
    """
    Queue background jobs for processing selected clients
    
    Args:
        client_names: List of client document names to process
    """
    if isinstance(client_names, str):
        # Convert string representation to list if passed as JSON string
        import json
        client_names = json.loads(client_names)
    
    # Queue each client for processing
    for client_name in client_names:
        # Get the client doc
        doc = frappe.get_doc("Client", client_name)
        
        
        # Queue the background job
        if not doc.disabled:
            frappe.enqueue(
                'fin_buddy.events.incometax_gov.process_single_client',
                queue='long',
                timeout=3000,
                client_name=client_name
            )
    
    return {"message": "Successfully queued selected clients for processing"}


def process_single_client(client_name):
    """Process a single client in the background
    
    Args:
        client_name: Name of the client document to process
    """
    try:
        # Get the client document
        doc = frappe.get_doc("Client", client_name)
        
        
        # Get credentials
        username = doc.username
        password = doc.get_password("password")
        
        if not username or not password:
            frappe.log_error(
                f"Missing credentials for client {client_name}",
                "Income Tax Portal Sync Error"
            )
            return
        
        # Setup Selenium and process
        driver = setup_driver(username, "incometax")
        
        start_time = time.time()
        
        try:
            if login_user(driver, username, password):
                # Check statuses
                worklist_status, demands = check_response_to_outstanding_demand_page(driver, client_name)
                # worklist_status, demands = None, None
                eproc_status = check_eproceedings_status(driver, client_name)

                # Logout the current user
                logout_user(driver)
                
            else:

                frappe.log_error('Error In User Login', 'Income Tax portal sync failed: Login failed')
        except Exception as e:
            error_msg = f"Error processing client {client_name}: {str(e)}"
            frappe.log_error(error_msg, "Income Tax Portal Processing Error")
        finally:
            # Calculate processing time
            doc.last_income_tax_sync = datetime.now()
            
            # Save the document
            doc.save()
            
            # Close the driver
            driver.quit()
    
    except Exception as e:
        error_msg = f"An error occurred in process_single_client for {client_name}: {str(e)}"
        frappe.log_error(error_msg, "Income Tax Portal Sync Error")




@frappe.whitelist()
def submit_response(client_name, notice_id):
    """
    Queue background jobs for processing a selected client.
    
    Args:
        client_name: Client document name to process.
    """
    if client_name:
        # Get the client doc
        doc = frappe.get_doc("Client", client_name)

        # Queue the background job
        if not doc.disabled:
            frappe.enqueue(
                'fin_buddy.events.incometax_gov.process_submit_response',
                queue='long',
                timeout=3000,
                client_name=client_name,
                notice_id=notice_id,
            )
        else:
            return {"message": f"Client {client_name} is disabled !!"}


        return {"message": f"Processing for {client_name} has started. Please wait while we submit your response !"}
    else:
        return {"message": "No client found to process"}


def process_submit_response(client_name, notice_id):
    try:
        print("processing start")
        doc = frappe.get_doc("Client", client_name)
        # Get credentials
        username = doc.username
        password = doc.get_password("password")
        
        if not username or not password:
            frappe.log_error(
                f"Missing credentials for client {client_name}",
                "Income Tax Portal Sync Error"
            )
            return
        
        # Setup Selenium and process
        driver = setup_driver(username, "incometax")
                
        try:
            login_user(driver, username, password)
            check_eproceedings_response_document(driver, client_name, notice_id)
            # handle_eproceedings_submit_response(driver, client_name, notice_id)
        except Exception as e:
            frappe.log_error("Login User Failed", str(e))
    
    except Exception as e:
        frappe.log_error("Login Error In Client", str(e))




def check_eproceedings_response_document(driver, client_name, notice_id):
    print("check_eproceedings_response_document start")
    try:
        driver.execute_script("document.body.scrollTop = 0; document.documentElement.scrollTop = 0;")

        # Allow time for scrolling
        time.sleep(1)
        # Click Pending Actions
        WebDriverWait(driver, 20).until(
            EC.visibility_of_element_located((By.ID, "Pending Actions"))
        ).click()

        # Click e-Proceedings
        e_proceedings_button = WebDriverWait(driver, 20).until(
            EC.visibility_of_element_located((By.XPATH, "//span[contains(text(), 'e-Proceedings')]"))
        )
        e_proceedings_button.click()
        
        # Wait for page to load
        time.sleep(3)
        
        page_source = driver.page_source
        soup = BeautifulSoup(page_source, 'html.parser')
        
        if "No e-Proceedings records are found" in soup.get_text():
            return "No Action Required"
        else:
            try:    
                handle_eproceedings_submit_response(driver, client_name, notice_id)
            except Exception as e:
                print(f"Failed to handle the dowloads in eproceeding Excel for user {client_name}: {str(e)}")

            return "Action Required"
    except Exception as e:
        return f"Error: {str(e)}"



def handle_eproceedings_submit_response(driver, client_name, notice_id):
    print("handle_eproceedings_submit_response start")

    def extract_page2_container_details(container_element):
        """Extract all details from a single notice container."""
        print("container")
        print(container_element)
        try:
            details = {}
            
            # Extract notice reference ID
            ref_id_elem = container_element.select_one("span:-soup-contains('Notice/ Communication Reference ID :')")
            details['notice_din'] = None

            if ref_id_elem:
                ref_id = ref_id_elem.find_next('span', class_='heading5')
                if ref_id:
                    details['notice_din'] = ref_id.text.strip()

            if not details['notice_din']:
                ref_id_elem = container_element.select_one("span:-soup-contains(' Document Identification Number (DIN) : ')")
                if ref_id_elem:
                    doc_id = ref_id_elem.find_next('span', class_='heading5')
                    if doc_id:
                        details['notice_din'] = doc_id.text.strip()

            
            # Extract notice number and type
            notice_elem = container_element.select_one("div.heading6")
            if notice_elem and notice_elem.text.strip().isdigit():
                details['notice_section'] = notice_elem.text.strip()
                details['notice_number'] = notice_elem.text.strip()
                

            # Step 1: First find the div that contains the "Document reference ID" text
            doc_ref_container = container_element.find("div", class_="dataHeading mt-1", text=lambda t: "Document reference ID" in t if t else False)

            if doc_ref_container:
                # Step 2: Get the parent div
                parent_div = doc_ref_container.parent
                
                if parent_div:
                    # Step 3: From the parent div, find the heading6 div
                    doc_id_elem = parent_div.find("div", class_="heading6 lineSpace")
                    
                    if doc_id_elem:
                        doc_id = doc_id_elem.get_text(strip=True)
                        details['notice_communication_reference_id'] = doc_id.strip()
                        print("✅ Extracted Document Reference ID:", doc_id)
                    else:
                        print("❌ No <div class='heading6 lineSpace'> found inside parent")
            
            
            # Extract description
            desc_elem = container_element.select_one("span:-soup-contains('Description :')")
            if desc_elem:
                description = desc_elem.find_next('span', class_='subtitle2')
                if description:
                    details['description'] = description.text.strip()
            
            # Extract dates
            issued_date_elem = container_element.select_one("span:-soup-contains('Issued On :')")
            if issued_date_elem:
                issued_date = issued_date_elem.find_next('span', class_='subtitle2')
                if issued_date:
                    notice_sent_date_value = issued_date.text.strip()
                    converted_date = MakeDateValid.get_frappe_date(notice_sent_date_value)
                    details['notice_sent_date'] = converted_date
                    
            due_date_elem = container_element.select_one("span:-soup-contains('Response Due Date :')")
            if due_date_elem:
                due_date = due_date_elem.find_next('span', class_='subtitle2')
                if due_date:
                    response_due_date_value = due_date.text.strip()
                    converted_date = MakeDateValid.get_frappe_date(response_due_date_value)
                    details['response_due_date'] = converted_date


            return details
        except Exception as e:
            print(f"Error extracting container details: {str(e)}")
            return {}


    def process_page2_buttons_old(driver, notice_id):
        doc = frappe.get_doc("E Proceeding", notice_id)
        doc.notice_communication_reference_id
        doc.notice_din

        # using them find the matching conatiner and then click on the submit button of that conatiner 

        """Helper function to process buttons on page 2 and create/update E Proceeding records"""
        try:
            # Extract header details first (common to all notices on this page)
            # header_details = extract_header_details(driver)
            button_name = "Submit Response"
                # Find View Response buttons
            buttons = WebDriverWait(driver, 20).until(
                    EC.presence_of_all_elements_located((By.XPATH, "//button[contains(@class, 'defaultButton primaryButton cardbutton') and contains(text(), 'Submit Response')]"))
                )

            if not buttons:
                print(f"No {button_name} buttons found")
                return 0, []

            print(f"Found {len(buttons)} {button_name} buttons")
            processed = 0
            created_or_updated_docs = []


            # Now there also on this page have multiple container elements we have to find that one conatiner element item with the matching notice_din or notice_communication_reference_id 

            for j in range(len(buttons)):
                try:
                    # Get page source for extraction before clicking button
                    soup = BeautifulSoup(driver.page_source, 'html.parser')
                    
                    # Re-find buttons to avoid stale elements
                    current_buttons = WebDriverWait(driver, 20).until(
                        EC.presence_of_all_elements_located((By.XPATH, 
                            "//button[contains(@class, 'defaultButton primaryButton cardbutton') and contains(text(), 'Submit Response')]"
                        ))
                    )
                    if j >= len(current_buttons):
                        continue
                    
                    button = current_buttons[j]
                    
                    # Find the container for this specific button
                    card_container = button.find_element(By.XPATH, "./ancestor::div[contains(@class, 'card-container')]")
                    card_html = card_container.get_attribute('outerHTML')
                    container_soup = BeautifulSoup(card_html, 'html.parser')
                    
                    # Extract container details
                    container_details = extract_page2_container_details(container_soup)
                    
                    # Merge header and container details
                    proceeding_details = {**container_details}
                    
                    print(f"Processing {button_name} button {j + 1}/{len(buttons)}")
                    
                    # Click button to go to third page
                    driver.execute_script("arguments[0].scrollIntoView({block: 'center'});", button)
                    time.sleep(2)
                    driver.execute_script("arguments[0].click();", button)
                    time.sleep(3)
                    
                   
                    
                except Exception as e:
                    print(f"Error processing {button_name} button {j + 1}: {str(e)}")
                    continue

            return processed, created_or_updated_docs

        except Exception as e:
            # print(f"Error in processing {button_type} buttons: {str(e)}")
            return 0, []

    def process_page2_buttons(driver, notice_id):
        try:
            doc = frappe.get_doc("E Proceeding", notice_id)
            print(f"Looking for container matching DIN: {doc.notice_din} or Reference ID: {doc.notice_communication_reference_id}")
            
            # Get all containers first
            containers = WebDriverWait(driver, 20).until(
                EC.presence_of_all_elements_located((By.CLASS_NAME, "card-container"))
            )
            
            matching_container = None
            matching_button = None
            
            # Iterate through containers to find match
            for container in containers:
                try:
                    # Get container HTML for BeautifulSoup parsing
                    container_html = container.get_attribute('outerHTML')
                    container_soup = BeautifulSoup(container_html, 'html.parser')
                    
                    # Extract details using existing function
                    container_details = extract_page2_container_details(container_soup)
                    
                    # Check for matches
                    if (doc.notice_din and container_details.get('notice_din') == doc.notice_din) or \
                    (doc.notice_communication_reference_id and 
                        container_details.get('notice_communication_reference_id') == doc.notice_communication_reference_id):
                        
                        print("Found matching container!")
                        matching_container = container
                        
                        # Find Submit Response button within this container
                        matching_button = container.find_element(
                            By.XPATH,
                            ".//button[contains(@class, 'defaultButton primaryButton cardbutton') and contains(text(), 'Submit Response')]"
                        )
                        break
                        
                except Exception as e:
                    print(f"Error processing container: {str(e)}")
                    continue
            
            if not matching_container or not matching_button:
                print("No matching container found")
                return 0, []
                
            try:
                # Click the Submit Response button for matching container
                print("Clicking Submit Response button for matching container")
                driver.execute_script("arguments[0].scrollIntoView({block: 'center'});", matching_button)
                time.sleep(2)
                driver.execute_script("arguments[0].click();", matching_button)
                time.sleep(3)

                continueBtn = WebDriverWait(driver, 10).until(
                    EC.visibility_of_element_located((By.ID, "instructionsPopupcontinueBtn"))
                )
                continueBtn.click()

                # Here you can add any additional processing needed after clicking Submit Response
                # For example, handling the next page or capturing results
                
                return 1, [notice_id]  # Return 1 for processed count
                
            except Exception as e:
                print(f"Error clicking Submit Response button: {str(e)}")
                return 0, []

        except Exception as e:
            print(f"Error in process_page2_buttons: {str(e)}")
            return 0, []
    
    
    
    
    try:
        doc = frappe.get_doc("E Proceeding", notice_id)
        processed_count = 0
        all_created_or_updated_docs = []
        
        # Wait for containers to be present
        containers = WebDriverWait(driver, 20).until(
            EC.presence_of_all_elements_located((By.CLASS_NAME, "card-container"))
        )
        
        matching_button = None
        
        # Iterate through containers to find matching criteria
        for container in containers:
            try:
                # Extract container details
                proceeding_name = container.find_element(By.XPATH, f".//span[contains(@class, 'heading5') and contains(text(), '{doc.proceeding_name}')]").text
                assessment_year = container.find_element(By.XPATH, ".//span[contains(text(), 'Assessment Year')]/following-sibling::span").text
                financial_year = container.find_element(By.XPATH, ".//span[contains(text(), 'Financial Year')]/following-sibling::span").text

                
                # Check if this container matches our criteria
                if (proceeding_name == doc.proceeding_name and 
                    assessment_year == doc.assessment_year and 
                    financial_year == doc.financial_year):
                    
                    # Find the View Notices/Orders button within this container
                    matching_button = container.find_element(
                        By.XPATH,
                        ".//button[contains(@class, 'defaultButton primaryButton') and contains(text(), 'View Notices/Orders')]"
                    )
                    break
                    
            except Exception as e:
                print(f"Error processing container: {str(e)}")
                continue
        
        if not matching_button:
            print("No matching container found")
            return "No matching items to process"
            
        # Process the matching container
        try:
            # Click View Notices/Orders to go to page 2
            driver.execute_script("arguments[0].scrollIntoView({block: 'center'});", matching_button)
            time.sleep(2)
            driver.execute_script("arguments[0].click();", matching_button)
            time.sleep(3)
            
            # Process all View Response buttons first
            response_count, response_docs = process_page2_buttons(driver, notice_id)

            time.sleep(2)
            # Move to top
            driver.execute_script("document.body.scrollTop = 0; document.documentElement.scrollTop = 0;")

                        # Find the element
            textarea = driver.find_element(By.ID, "inputcomment")

            # Clear any existing text first
            textarea.clear()

            # Scroll the element into view
            driver.execute_script("arguments[0].scrollIntoView(true);", textarea)
            time.sleep(1)  # Give it a moment to scroll

            # Set the value using JavaScript
            driver.execute_script("arguments[0].value = arguments[1];", textarea, doc.response_message)

            
        except Exception as e:
            print(f"Error processing matching container: {str(e)}")
            return f"Error: {str(e)}"
            
        # Return summary of processing
        return f"Downloads Completed - Processed {processed_count} items, Created/Updated {len(set(all_created_or_updated_docs))} E Proceeding records"
        
    except Exception as e:
        print(f"Error in handle_eproceedings_downloads: {str(e)}")
        return f"Error: {str(e)}"






# OPEN AI and response generation

@frappe.whitelist()
def fetch_response_from_gpt(doctype,docname, is_view_data_before_response_generation=False):

    # Get the document from the Doctype
    doc = frappe.get_doc(doctype, docname)

    settings = frappe.get_single("FinBuddy Settings")

    notice_text = "no notice attached for now."
    # Ensure the attachment exists
    if not settings.give_response_without_notice and not doc.notice_letter :
        frappe.throw("No Notice PDF attachment found.")
    elif doc.notice_letter:
            notice_text_doc = is_present_in_notice_text(doc.notice_letter)
            if notice_text_doc:
                notice_text = notice_text_doc.notice_text
            else:
                notice_text = extract_document_text(doc.notice_letter)
                try:
                    new_nt_doc = frappe.get_doc({
                        "doctype": "Notice Text",
                        # "doctype_name": "E Proceeding",
                        # "doc_name": docname,
                        "notice": doc.notice_letter,
                        "notice_text": notice_text
                    })

                    new_nt_doc.insert()
                    frappe.db.commit()
                except Exception as e:
                    frappe.log_error("Notice Text", f"Can't Create new notice text: \n {str(e)}")
                    frappe.throw("Not able to create the Notice Text")

                
    


    # User's query (you can get this from another field if needed)
    user_input = doc.user_input or "no user input."
    today = datetime.today()
    user_input += f"And Today date is {today.strftime('%d-%m-%Y')}. Other dates are notice dates like sent or other."
    # user_input += f"\n Proceeding Name is {doc.proceeding_name} and Financial Year is {doc.financial_year} and Document Identification Number is {doc.notice_din} and Notice Section is {doc.notice_section} and Assessment Year is {doc.assessment_year} and Document reference ID is {doc.notice_communication_reference_id}. If some value is not present then ignore that values."

    
    other_documents_text = ""
    for document in doc.other_documents:
        notice_text_doc = is_present_in_notice_text(document.file)
        if notice_text_doc:
                document_text = notice_text_doc.notice_text
        else:
                document_text = extract_document_text(document.file)
                try:
                    new_nt_doc = frappe.get_doc({
                        "doctype": "Notice Text",
                        # "doctype_name": "E Proceeding",
                        # "doc_name": docname,
                        "notice": document.file,
                        "notice_text": document_text
                    })

                    new_nt_doc.insert()
                    frappe.db.commit()
                except Exception as e:
                    frappe.log_error("Notice Text", f"Can't Create new notice text: \n {str(e)}")
                    frappe.throw("Not able to create the Notice Text")

        print(f"\n\n {document_text} \n\n")
        other_documents_text += f"Document Content: \n {document_text} \n"

    if not other_documents_text:
        other_documents_text = "no other documents attached."

    mask_this_data = doc.mask_this_data or ""

    if not is_view_data_before_response_generation:
        # Query OpenAI
        response = query_openai(user_input, notice_text, other_documents_text, mask_this_data, doc)

        # Save response in the response_message field
        doc.response_message = response
        doc.save()
        frappe.db.commit()

        return response  # Return the response to frontend (optional)
    else:
        return view_data_before_response_generation(user_input, notice_text, other_documents_text, mask_this_data)

def extract_pdf_text_old(pdf_path):
    """Extracts and cleans text from a given PDF file."""
    with open(pdf_path, 'rb') as file:
        reader = PyPDF2.PdfReader(file)
        text = []
        
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                # Normalize spaces and remove excessive newlines
                cleaned_text = re.sub(r'\s+', ' ', page_text).strip()
                text.append(cleaned_text)

    return ' '.join(text).strip()  # Join all pages into a single string



# def is_present_in_notice_text(doctype, docname, file_url):
def is_present_in_notice_text(file_url):
    existing_records = frappe.get_all("Notice Text",
        # filters = {"doctype_name": doctype, "doc_name": docname, "notice": file_url},
        filters = {"notice": file_url},
        fields=["name"]
    )

    if existing_records:
        return frappe.get_doc("Notice Text", existing_records[0]['name'])


def extract_document_text(file_url):
    """
    Extracts and cleans text from documents stored in Frappe's File Doctype.
    Supports PDF and DOCX formats.
    """
    # Fetch the file document from the database
    file_doc = frappe.get_doc("File", {"file_url": file_url})
    
    if not file_doc:
        raise FileNotFoundError(f"File not found in Frappe File Doctype: {file_url}")
    
    # Get the file content as bytes
    file_content = file_doc.get_content()
    
    if not file_content:
        raise ValueError(f"Failed to retrieve content for {file_url}")
    
    # Determine file type from extension
    file_extension = Path(file_url).suffix.lower()
    
    try:
        if file_extension == '.pdf':
            return _extract_pdf_text(file_content)
        elif file_extension in ['.docx', '.doc']:
            return _extract_docx_text(file_content)
        else:
            raise ValueError(f"Unsupported file format: {file_extension}")
    except Exception as e:
        raise ValueError(f"Error processing {file_url}: {str(e)}")

def _extract_pdf_text(file_content):
    """Helper function to extract text from PDF content."""
    text = []
    pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_content))
    
    for page in pdf_reader.pages:
        page_text = page.extract_text()
        if page_text:
            cleaned_text = re.sub(r'\s+', ' ', page_text).strip()
            text.append(cleaned_text)
    
    return ' '.join(text).strip()

def _extract_docx_text(file_content):
    """Helper function to extract text from DOCX content."""
    text = []
    doc = Document(io.BytesIO(file_content))
    
    # Extract text from paragraphs
    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            cleaned_text = re.sub(r'\s+', ' ', paragraph.text).strip()
            text.append(cleaned_text)
    
    # Extract text from tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    cleaned_text = re.sub(r'\s+', ' ', cell.text).strip()
                    text.append(cleaned_text)
    
    return ' '.join(text).strip()

def query_openai_old(user_input, notice_text, other_documents_text):
    """Call OpenAI API to generate response."""

    # Get GPT model from settings
    settings = frappe.get_single("FinBuddy Settings")
    model = settings.model_name or "gpt-3.5-turbo"
    max_token = settings.max_token or 500
    temperature = settings.temperature or 0.2
    frequency_penalty = settings.frequency_penalty or 0.3
    system_role = settings.system_role or "give clear and structured response."


    openai.api_key = settings.get_password('openai_key') # key

    combined_user_content = f"Notice File Content: {notice_text} \n Additional Documents: {other_documents_text} \n User Input: {user_input}"


    response = openai.ChatCompletion.create(
        model=model,
        messages=[
            {"role": "system", "content": system_role},
            {"role": "user", "content": combined_user_content.strip()},
        ],
        max_tokens=max_token,
        temperature=temperature,
        frequency_penalty=frequency_penalty
    )

    return response["choices"][0]["message"]["content"].strip()


def query_openai_fine(user_input, notice_text, other_documents_text):
    """Call OpenAI API to generate response with content masking."""
    
    # Initialize settings and masker
    settings = frappe.get_single("FinBuddy Settings")
    masker = ContentMasker()
    
    # Combine all text that needs to be masked
    combined_user_content = f"Notice File Content: {notice_text} \n Additional Documents: {other_documents_text} \n User Input: {user_input}"
    
    # Mask the content
    masked_content = masker.mask_content(combined_user_content)
    
    # Save mapping to FinBuddy DocType
    mapping_data = {
        'mapping': masker.mapping,
        'reverse_mapping': masker.reverse_mapping,
        'timestamp': str(datetime.now())
    }
    
    settings.mapping_content = json.dumps(mapping_data)
    settings.save()

    # Get OpenAI settings
    model = settings.model_name or "gpt-3.5-turbo"
    max_token = settings.max_token or 500
    temperature = settings.temperature or 0.2
    frequency_penalty = settings.frequency_penalty or 0.3
    system_role = settings.system_role or "give clear and structured response."
    openai.api_key = settings.get_password('openai_key')

    # Make API call with masked content
    response = openai.ChatCompletion.create(
        model=model,
        messages=[
            {"role": "system", "content": system_role},
            {"role": "user", "content": masked_content.strip()},
        ],
        max_tokens=max_token,
        temperature=temperature,
        frequency_penalty=frequency_penalty
    )

    # Get the response
    masked_response = response["choices"][0]["message"]["content"].strip()

    # Load mapping from FinBuddy DocType
    stored_mapping = json.loads(settings.mapping_content)
    masker.mapping = stored_mapping['mapping']
    masker.reverse_mapping = stored_mapping['reverse_mapping']

    # Unmask the response
    unmasked_response = masker.unmask_content(masked_response)

    return unmasked_response





# def query_openai(user_input: str, notice_text: str, other_documents_text: str) -> str:
#     """Process content through OpenAI with masking/unmasking"""
#     try:
#         # Initialize
#         settings = frappe.get_single("FinBuddy Settings")
#         system_role = settings.system_role or "give me proper response."
#         masker = ContentMasker()
#         log_filepath = None
        
#         # Log initial inputs
#         initial_content = {
#             "user_input": user_input,
#             "notice_text": notice_text,
#             "other_documents_text": other_documents_text
#         }
#         log_filepath = write_to_content_file(initial_content, "1. INITIAL INPUT")
        
#         # Combine content
#         combined_user_content = f"""Notice File Content: {notice_text} \nAdditional Documents :{other_documents_text} \nUser Input: {user_input}"""
        
#         write_to_content_file(combined_user_content, "2. COMBINED CONTENT", log_filepath)
        
#         # Mask content
#         masked_content = masker.mask_content(combined_user_content)
#         write_to_content_file({
#             "masked_content": masked_content,
#             "mapping": masker.mapping,
#             "reverse_mapping": masker.reverse_mapping
#         }, "3. MASKED CONTENT AND MAPPING", log_filepath)
        
#         # Save mapping to settings
#         mapping_data = {
#             'mapping': masker.mapping,
#             'reverse_mapping': masker.reverse_mapping,
#             'timestamp': str(datetime.now()),
#             'content_log': masker.content_log
#         }
#         settings.mapping_content = json.dumps(mapping_data)
#         settings.save()
                
#         # Get OpenAI settings and make request
#         openai.api_key = settings.get_password('openai_key')
#         openai_request = {
#             "model": settings.model_name or "gpt-3.5-turbo",
#             "max_tokens": settings.max_token or 500,
#             "temperature": settings.temperature or 0.2,
#             "frequency_penalty": settings.frequency_penalty or 0.3,
#             "system_role": system_role
#         }
#         write_to_content_file(openai_request, "4. OPENAI REQUEST", log_filepath)
        
#         # Make API call
#         response = openai.ChatCompletion.create(
#             model=openai_request["model"],
#             messages=[
#                 {"role": "system", "content": system_role},
#                 {"role": "user", "content": masked_content.strip()}
#             ],
#             max_tokens=openai_request["max_tokens"],
#             temperature=openai_request["temperature"],
#             frequency_penalty=openai_request["frequency_penalty"]
#         )
        
#         # Process response
#         masked_response = response["choices"][0]["message"]["content"].strip()
#         write_to_content_file(masked_response, "5. MASKED RESPONSE FROM GPT", log_filepath)
        
#         # Unmask response
#         stored_mapping = json.loads(settings.mapping_content)
#         masker.mapping = stored_mapping['mapping']
#         masker.reverse_mapping = stored_mapping['reverse_mapping']
#         unmasked_response = masker.unmask_content(masked_response)
        
#         write_to_content_file(unmasked_response, "6. FINAL UNMASKED RESPONSE", log_filepath)
        
#         return unmasked_response
        
#     except Exception as e:
#         error_msg = f"Error in query_openai: {str(e)}"
#         if log_filepath:
#             write_to_content_file({"error": error_msg}, "ERROR", log_filepath)
#         frappe.log_error(error_msg)
#         raise


# def write_to_content_file(content, stage, filepath=None):
#     """Write content to tracking file with stage information"""
#     try:
#         # Create logs directory if it doesn't exist
#         log_dir = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'content_logs')
#         os.makedirs(log_dir, exist_ok=True)
        
#         # Use provided filepath or create new one
#         if not filepath:
#             filename = f'content_tracking_{datetime.now().strftime("%Y%m%d_%H%M%S")}.txt'
#             filepath = os.path.join(log_dir, filename)
        
#         with open(filepath, 'a', encoding='utf-8') as f:
#             f.write(f"\n{'='*50}\n")
#             f.write(f"STAGE: {stage}\n")
#             f.write(f"TIMESTAMP: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
            
#             if isinstance(content, dict):
#                 f.write("CONTENT:\n")
#                 f.write(json.dumps(content, indent=2, ensure_ascii=False))
#             else:
#                 f.write("CONTENT:\n")
#                 f.write(str(content))
            
#             f.write(f"\n{'='*50}\n")
        
#         return filepath
    
#     except Exception as e:
#         frappe.log_error(f"Error writing to content tracking file: {str(e)}")
#         return ""



def view_data_before_response_generation(user_input: str, notice_text: str, other_documents_text: str, mask_this_data: str) -> str:
    try:
        # Initialize
        masker = DataMaskingUtil()
        
        # Log initial inputs
        initial_content = f"User Input:\n{user_input}\n\nNotice Data:\n{notice_text}\n\nOther Document Text:\n{other_documents_text}\n\nMask This Data:\n{mask_this_data}\n"
        # log_filepath = write_to_content_file(initial_content, "1. INITIAL INPUT")
        
        # Combine content
        combined_user_content = f"""Notice File Content:\n{notice_text}\n\nAdditional Documents:\n{other_documents_text}\n\nUser Input:\n{user_input}\n"""
        
        # Process mask_this_data into list
        entities_to_mask = [
            item.strip() 
            for item in mask_this_data.split(',') 
            if item.strip()  # Only include non-empty items
        ]
        
        # First mask the explicitly provided entities
        masked_content = masker.mask_text(combined_user_content, entities_to_mask)
        
        # Then mask other sensitive data (PAN, dates, etc.)
        masked_content = masker.mask_text(masked_content)  # Second pass without entities list
        

        masked_content_and_map = f"Masked Content Which We Will Send To AI For Response Generation\n\n{masked_content}\n\nMapping:\n{masker.mapping}\n\nEntity To Mask:\n{entities_to_mask}\n"
        
        return {"data": masked_content_and_map}
        
    except Exception as e:
        frappe.log_error("problem in view data before response generation", str(e))
        return {"data": "Not abel to show Masked Content and Mapping."}


def query_openai(user_input: str, notice_text: str, other_documents_text: str, mask_this_data: str, current_doc) -> str:
    """
    Process content through OpenAI with masking/unmasking
    
    Args:
        user_input: User's question or input
        notice_text: Main notice content
        other_documents_text: Additional document content
        mask_this_data: Comma-separated text of items to mask (e.g., "Nav Bharat, Hisar, John Doe")
    """
    try:
        # Initialize
        settings = frappe.get_single("FinBuddy Settings")
        system_role = settings.system_role or "give me proper response."
        masker = DataMaskingUtil()
        log_filepath = None
        
        # Log initial inputs
        initial_content = f"User Input:\n{user_input}\n\nNotice Data:\n{notice_text}\n\nOther Document Text:\n{other_documents_text}\n\nMask This Data:\n{mask_this_data}\n"
        # log_filepath = write_to_content_file(initial_content, "1. INITIAL INPUT")

        doc_info = {
            'client_name': current_doc.client,
            'doctype_name': current_doc.doctype,
            'doc_name': current_doc.name,
        }
        response_doc_log = write_content_generation_logs(initial_content, 'initial_content', None, doc_info)

        print("respne doc log")
        print(response_doc_log)
        
        # Combine content
        combined_user_content = f"""Notice File Content:\n{notice_text}\n\nAdditional Documents:\n{other_documents_text}\n\nUser Input:\n{user_input}\n"""
        
        # write_to_content_file(combined_user_content, "2. COMBINED CONTENT", log_filepath)
        response_doc_log = write_content_generation_logs(combined_user_content, 'combined_content', response_doc_log, doc_info)
        # Process mask_this_data into list
        entities_to_mask = [
            item.strip() 
            for item in mask_this_data.split(',') 
            if item.strip()  # Only include non-empty items
        ]
        
        # First mask the explicitly provided entities
        masked_content = masker.mask_text(combined_user_content, entities_to_mask)
        
        # Then mask other sensitive data (PAN, dates, etc.)
        masked_content = masker.mask_text(masked_content)  # Second pass without entities list
        
        # write_to_content_file({
        #     "masked_content": masked_content,
        #     "mapping": masker.mapping,
        #     "entities_masked": entities_to_mask
        # }, "3. MASKED CONTENT AND MAPPING", log_filepath)
        
        masked_content_and_map = f"Masked Content:\n{masked_content}\n\nMapping:\n{masker.mapping}\n\nEntity To Mask:\n{entities_to_mask}\n"

        response_doc_log = write_content_generation_logs(masked_content_and_map, 'masked_content', response_doc_log, doc_info)
        
        # Save mapping to settings
        mapping_data = {
            'mapping': masker.mapping,
            'timestamp': str(datetime.now()),
            'entities_masked': entities_to_mask
        }
        settings.mapping_content = json.dumps(mapping_data)
        settings.save()
        
        # Get OpenAI settings and make request
        openai.api_key = settings.get_password('openai_key')
        openai_request = {
            "model": settings.model_name or "gpt-3.5-turbo",
            "max_tokens": settings.max_token or 500,
            "temperature": settings.temperature or 0.2,
            "frequency_penalty": settings.frequency_penalty or 0.3,
            "system_role": system_role
        }
        # write_to_content_file(openai_request, "4. OPENAI REQUEST", log_filepath)
        
        # Make API call
        response = openai.ChatCompletion.create(
            model=openai_request["model"],
            messages=[
                {"role": "system", "content": system_role},
                {"role": "user", "content": masked_content.strip()}
            ],
            max_tokens=openai_request["max_tokens"],
            temperature=openai_request["temperature"],
            frequency_penalty=openai_request["frequency_penalty"]
        )
        
        # Process response
        masked_response = response["choices"][0]["message"]["content"].strip()
        # write_to_content_file(masked_response, "5. MASKED RESPONSE FROM GPT", log_filepath)

        response_doc_log = write_content_generation_logs(masked_response, 'masked_response', response_doc_log, doc_info)
        
        # Unmask response
        stored_mapping = json.loads(settings.mapping_content)
        masker.mapping = stored_mapping['mapping']
        unmasked_response = masker.restore_text(masked_response)
        
        # write_to_content_file(unmasked_response, "6. FINAL UNMASKED RESPONSE", log_filepath)
        response_doc_log = write_content_generation_logs(unmasked_response, 'final_response', response_doc_log, doc_info)


        # print("before insert")
        # print(response_doc_log.as_dict())

        response_doc_log.insert()
        frappe.db.commit()
        
        return unmasked_response
        
    except Exception as e:
        error_msg = f"Error in query_openai: {str(e)}"
        if log_filepath:
            write_to_content_file({"error": error_msg}, "ERROR", log_filepath)
        frappe.log_error(error_msg)
        raise

def write_to_content_file(content: Any, stage: str, filepath: Optional[str] = None) -> str:
    """Write content to tracking file with stage information"""
    try:
        # Create logs directory if it doesn't exist
        log_dir = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'content_logs')
        os.makedirs(log_dir, exist_ok=True)
        
        # Use provided filepath or create new one
        if not filepath:
            filename = f'content_tracking_{datetime.now().strftime("%Y%m%d_%H%M%S")}.txt'
            filepath = os.path.join(log_dir, filename)
        
        with open(filepath, 'a', encoding='utf-8') as f:
            f.write(f"\n{'='*50}\n")
            f.write(f"STAGE: {stage}\n")
            f.write(f"TIMESTAMP: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
            
            if isinstance(content, dict):
                f.write("CONTENT:\n")
                f.write(json.dumps(content, indent=2, ensure_ascii=False))
            else:
                f.write("CONTENT:\n")
                f.write(str(content))
            
            f.write(f"\n{'='*50}\n")
        
        return filepath
        
    except Exception as e:
        frappe.log_error(f"Error writing to content tracking file: {str(e)}")
        return ""




def write_content_generation_logs(content, stage_name, doc_name=None, doc_info=None):
    try:
        print("stage_name", stage_name)
        print("doc_name", doc_name)
        print("doc_info", doc_info)
        if not doc_name:
            doc_name = frappe.get_doc({
                'doctype': "Generate Response Log",
                'client': doc_info['client_name'],
                'doctype_name': doc_info['doctype_name'],
                'doc_name': doc_info['doc_name'],
            })

            print("my_doc")

        doc_name.set(stage_name, content)
        
        return doc_name
        
    except Exception as e:
        frappe.log_error(f"Error creating to response generation log tracking: {str(e)}")
        return None